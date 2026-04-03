"""DOCX builder — assembles manuscript sections into a formatted Word document."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from publicaition.orchestrator.state import Draft

# Section display labels — maps section_type to readable heading
SECTION_LABELS: dict[str, str] = {
    "abstract":      "Abstract",
    "introduction":  "Introduction",
    "methods":       "Methods",
    "results":       "Results",
    "discussion":    "Discussion",
    "conclusion":    "Conclusion",
    "reference_list": "Citation Annotations & References",
}


def build_docx(
    output_path: Path,
    project_name: str,
    journal: str,
    sections: list[Draft],
    reference_list: Draft | None,
    pls: Draft | None,
) -> Path:
    doc = Document()
    _configure_page(doc)
    _add_title_page(doc, project_name, journal)
    _add_page_break(doc)

    for draft in sections:
        _add_section(doc, draft, SECTION_LABELS.get(draft.section_type, draft.section_type.title()))
        _add_page_break(doc)

    if reference_list and reference_list.text.strip():
        _add_section(doc, reference_list, "References")
        _add_page_break(doc)

    if pls and pls.text.strip():
        _add_section(doc, pls, "Plain Language Summary")

    doc.save(str(output_path))
    return output_path


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1.0)
    section.top_margin  = section.bottom_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = Pt(24)   # double-spaced


def _add_title_page(doc: Document, project_name: str, journal: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(project_name)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()  # spacer

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Target Journal: {journal}\n")
    meta.add_run(f"Generated: {date.today().isoformat()}\n")
    meta.add_run("Confidential — For Review Only")


def _add_section(doc: Document, draft: Draft, label: str) -> None:
    heading = doc.add_heading(label, level=1)
    heading.style.font.name = "Times New Roman"
    heading.style.font.size = Pt(14)

    wc_note = doc.add_paragraph()
    wc_note.add_run(f"[{draft.word_count} words]").italic = True
    wc_note.paragraph_format.space_after = Pt(6)

    for paragraph_text in draft.text.split("\n\n"):
        paragraph_text = paragraph_text.strip()
        if not paragraph_text:
            continue
        p = doc.add_paragraph(paragraph_text)
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(0)


def _add_page_break(doc: Document) -> None:
    doc.add_page_break()
